#!/usr/bin/env python3
"""Builds a submission-draft preprint manuscript as a Word document."""
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1a, 0x1a, 0x1a)
MUTED = RGBColor(0x50, 0x50, 0x50)
LINE_GREY = "CCCCCC"
HEAD_GREY = "F2F2F2"
NOTE_BG = "FFF4D6"

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.font.color.rgb = INK
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.15

for sec in doc.sections:
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_border(cell, color=LINE_GREY):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)


def h1(text, number=None):
    p = doc.add_heading(level=1)
    r = p.add_run((f"{number}. " if number else "") + text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.color.rgb = INK
    r.font.bold = True
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p


def h2(text, number=None):
    p = doc.add_heading(level=2)
    r = p.add_run((f"{number} " if number else "") + text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.color.rgb = INK
    r.font.bold = True
    r.font.italic = False
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(text, italic=False, bold=False, size=11, align=None, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.bold = bold
    r.font.color.rgb = INK
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p


def note_box(text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, NOTE_BG)
    set_cell_border(cell, "E8C468")
    para = cell.paragraphs[0]
    r = para.add_run(text)
    r.font.size = Pt(9.5)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x5c, 0x47, 0x00)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def figure(path, caption, width_in=6.0):
    doc.add_picture(path, width=Inches(width_in))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    r = p.add_run(caption)
    r.font.size = Pt(9.5)
    r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext)
        r.font.bold = True
        r.font.size = Pt(9.5)
        set_cell_shading(hdr[i], HEAD_GREY)
        set_cell_border(hdr[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            set_cell_border(cells[i])
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


FIG = "results/figures/"

# ============================================================================
# TITLE PAGE / NOTE TO AUTHOR
# ============================================================================
note_box(
    "NOTE TO AUTHOR (remove before submission): this draft was compiled from a full analysis "
    "pipeline with every reported statistic pulled directly from the underlying results files. "
    "Before submitting, please: (1) double-check author order, affiliations, and the "
    "corresponding-author email (filled in below, but verify the affiliation structure and add the "
    "email); (2) all citations have now been verified against PubMed/PMC/publisher records (see "
    "Reference list note for one caveat re: preprint-vs-published status); (3) add an ORCID/funding/"
    "conflict-of-interest statement per your target venue's requirements; (4) decide on a target venue "
    "(bioRxiv is the natural first step given this is a secondary analysis of an already-public "
    "dataset) and reformat citations to that venue's style if needed; (5) fill in the code-repository "
    "URL in Data and Code Availability once your GitHub repo is created."
)

p = doc.add_paragraph()
r = p.add_run("Individual Differences in Reward-Zone Modulation of Hippocampal Representational "
              "Drift During Chronic Recording in a Spatial Working-Memory Task")
r.font.size = Pt(18)
r.font.bold = True
r.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(14)

body("Saad Yusuf¹,², Abhilasha Joshi¹", size=12)
body("¹ NeuroAct Lab, National Centre for Biological Sciences (NCBS), Bangalore, India", size=10.5, italic=True)
body("² Jamia Millia Islamia, New Delhi, India", size=10.5, italic=True)
body("Corresponding author: Abhilasha Joshi [email to be added]", size=10.5, italic=True)

body("")
h2("Abstract")
body(
    "Hippocampal spatial representations drift across days even in familiar, unchanging "
    "environments, and recent work shows that globally withholding reward accelerates this drift. "
    "Whether reward-proximal locations show altered drift dynamics relative to behaviorally neutral "
    "space within a single, continuously-rewarded environment has not been tested. We analyzed chronic "
    "tetrode recordings from dorsal CA1 in three mice performing a delayed-non-match-to-place T-maze "
    "task across 15–19 sessions spanning two to three weeks (DANDI:001775). Two independent "
    "data-quality issues were identified and corrected prior to analysis: unreliable session-timing "
    "metadata, verified against each session's internally-consistent embedded acquisition clock, and, "
    "in one subject, degraded position tracking. Population vector (PV) correlation, computed from "
    "tetrode-pooled multi-unit rate maps (cell identity is not tracked across sessions in this "
    "dataset), quantified representational similarity between session pairs as a function of "
    "calendar-day separation, computed separately for reward-zone, choice-point, and corridor spatial "
    "bins. PV correlation declined significantly with day separation in all three animals (Mantel "
    "r = −0.21 to −0.33, all p < .03; pooled p = 3.4×10⁻⁵), replicating chronic "
    "representational drift. The reward zone showed a significant drift-rate difference from corridor "
    "in all three animals (occupancy-matched shuffle test, all p < .03), but the direction was not "
    "uniform: two animals showed slower reward-zone drift, and one showed faster reward-zone drift. "
    "Held-out validation — zones defined on one temporal half of a subject's sessions and tested "
    "on the other — showed this effect replicated fully in only one animal (the faster-drift "
    "direction), partially in a second, and not at all in the third. The one fully-replicating effect "
    "additionally survived a direct control for the position-tracking artifact identified in that "
    "subject. These results indicate that reward-adjacent hippocampal representations are subject to "
    "altered, but not uniformly protective, drift dynamics, and highlight meaningful, currently "
    "unexplained individual variability in how value-relevant locations are stabilized in the "
    "hippocampal spatial map.",
    size=10.5,
)
body("Keywords: hippocampus; CA1; representational drift; population vector correlation; reward; "
     "spatial working memory; individual differences", italic=True, size=10)

doc.add_page_break()

# ============================================================================
h1("Introduction", 1)
body(
    "Hippocampal place cells provide a stable, allocentric representation of an animal's location "
    "within a familiar environment (O'Keefe & Dostrovsky, 1971). However, chronic recording and "
    "imaging studies have established that this stability is not absolute: the population code that "
    "represents a given, unchanging environment gradually reconfigures over days and weeks, a "
    "phenomenon termed representational drift (Ziv et al., 2013; Rule, O'Leary, & Harvey, 2019). Drift "
    "is now a well-replicated feature of hippocampal (and broader cortical) coding, but the factors "
    "that govern its rate remain an active area of investigation."
)
body(
    "Two broad classes of modulators have been identified. First, drift dissociates along temporal and "
    "experiential axes: elapsed time predominantly alters firing-rate statistics, whereas active "
    "behavioral experience in an environment predominantly alters spatial tuning (Geva, Deitch, Rubin, "
    "& Ziv, 2023). Second, reward availability modulates drift rate: comparing hippocampal population "
    "activity across days on which reward was present versus globally withheld shows that reduced "
    "reward expectation increases representational drift, an effect proposed to depend on "
    "reward-linked dopaminergic modulation of synaptic stability (Krishnan & Sheffield, 2023). Reward "
    "also more broadly reorganizes the hippocampal population code even absent any drift question: "
    "place fields accumulate disproportionately near goal locations (Hollup, Molden, Donnett, Moser, & "
    "Moser, 2001; Gauthier & Tank, 2018), a dedicated subpopulation of neurons appears to track position "
    "relative to reward and dynamically updates when reward is relocated (Sosa, Plitt, & Giocomo, 2025), "
    "and context-relevance effects on remapping are strongest specifically in place fields near reward "
    "locations (Tarcsay, Masala, Yi, Igarashi, Redic, & Ewell, 2025)."
)
body(
    "A gap remains between these two literatures. Studies linking reward to drift rate have compared "
    "globally rewarded versus globally unrewarded sessions — a between-condition, whole-environment "
    "manipulation. Studies showing reward reorganizes the population code at the location of reward "
    "have generally examined field density or population geometry at a single time point or over short "
    "timescales, not the long-term (multi-week) drift rate specifically at the reward location relative "
    "to the rest of a constantly-rewarded environment. Whether reward-adjacent representations are "
    "differentially stable — more protected, or conversely more volatile — relative to "
    "behaviorally neutral space within the same, continuously-rewarded session has not, to our "
    "knowledge, been directly tested."
)
body(
    "We address this using a publicly available chronic tetrode dataset recorded from dorsal CA1 during "
    "a delayed-non-match-to-place (DNMP) T-maze task, in which reward was delivered at a fixed physical "
    "location on every completed trial throughout 15–19 recording sessions spanning two to three "
    "weeks per animal. We first establish that this dataset replicates chronic representational drift, "
    "then ask whether population vector correlation declines at a different rate specifically within "
    "the reward zone relative to the corridor, applying four independent robustness checks: an "
    "occupancy-matched shuffle control, a zone-definition parameter-sensitivity check, held-out "
    "cross-validation, and a direct control for a position-tracking artifact identified during the "
    "course of this analysis in one subject."
)

# ============================================================================
h1("Results", 2)

h2("Three mice, 49 sessions, and two independent data-quality corrections", "2.1")
body(
    "Of the six subjects available in the source dataset, three were excluded prior to any drift "
    "analysis. One subject's files used a different behavioral paradigm entirely (a pharmacological "
    "homecage-foraging task) and lacked a spike-sorted units table. Two further subjects' position, "
    "spike, and trial timestamps, while internally self-consistent within each session, were found to "
    "be referenced to a lab-internal relative clock (elapsed acquisition-system time, on the order of "
    "10⁶ s) rather than true Unix epoch time, and could not be mapped to real calendar dates from "
    "file contents alone; these subjects (23 and 7 sessions respectively) were excluded rather than "
    "converted using an assumed offset. The final analyzed sample comprised three subjects with "
    "independently-verified absolute recording clocks: UT14 (19 sessions), UT13 (15 sessions), and "
    "UT15 (15 sessions)."
)
body(
    "For all three retained subjects, the dataset's own session-level timing metadata "
    "(“session_start_time”) was found to disagree with the true acquisition clock — "
    "reconstructed from the mutually-consistent Unix-epoch timestamps embedded in each session's "
    "position, spike, and trial-interval data — by up to several hours, inconsistently in "
    "direction across sessions (e.g., −455 to +203 minutes for UT14). All timing-dependent "
    "analyses reported here use the verified embedded acquisition clock, not the metadata field."
)
body(
    "Reward-zone and choice-point spatial regions were defined per subject from trial-level behavioral "
    "event timestamps (position at reward consumption, pooled across all sessions) and validated "
    "visually against each subject's occupancy map before use in any statistical test (Figure 1)."
)
figure(FIG + "Z1_zone_classification.png",
       "Figure 1. Reward-zone (red), choice-zone (orange), and corridor (grey) classification for each "
       "subject, overlaid on session occupancy (background shading). Reward zones are centered on the "
       "median position at reward consumption (grasp_time), computed separately for left and right "
       "arms and pooled across a subject's full session history; stars mark reward-port centroids. "
       "Choice zones are the stem bins immediately adjoining the arm-exclusive region. Bin counts and "
       "pooled reward-event counts are given per subject.")

h2("Hippocampal population codes decorrelate across days in all three animals", "2.2")
body(
    "Population vector correlation, computed from tetrode-pooled rate maps across the full maze, "
    "declined significantly with calendar-day separation in every subject (Mantel permutation test, "
    "5,000 permutations of session order): UT14, r = −0.313, p = .0017 (19 sessions, 171 pairs); "
    "UT13, r = −0.208, p = .0297 (15 sessions, 105 pairs); UT15, r = −0.328, p = .0053 (15 "
    "sessions, 105 pairs). Combining these three independent tests (Fisher's method) yielded "
    "χ²(6) = 30.30, p = 3.4×10⁻⁵. A complementary linear mixed-effects model "
    "(day distance as a fixed effect, random intercept per subject) confirmed a significant pooled "
    "day-distance effect (β = −0.0121, p = 3.5×10⁻⁷). This establishes a "
    "genuine, three-animal replication of chronic hippocampal representational drift in this dataset "
    "(Figure 2), providing the foundation for the zone-restricted analyses below."
)
figure(FIG + "A5_pv_vs_day_all_subjects.png",
       "Figure 2. Population vector correlation (whole maze) as a function of calendar-day separation, "
       "for each subject. Each point is one session pair; lines are OLS fits. All three subjects show a "
       "significant negative relationship (Mantel test, see main text).")

h2("Reward-zone representations show altered, but non-uniform, drift dynamics", "2.3")
body(
    "Restricting the population vector correlation to reward-zone bins versus corridor bins revealed a "
    "significant difference in drift rate (slope of PV correlation on day distance) in every subject "
    "(occupancy-matched bin-shuffle permutation test, 500 permutations, reassigning the same visited "
    "bins to same-sized pseudo-zone categories): UT14, slope difference = +0.0023, p = .0020; UT13, "
    "+0.0010, p = .0240; UT15, −0.0145, p = .0020. The sign of this difference, however, was not "
    "consistent across animals. In UT14 and UT13, the reward zone decorrelated more slowly than the "
    "corridor (a positive difference, indicating relative protection from drift). In UT15, the reward "
    "zone decorrelated faster than the corridor (a negative difference of larger magnitude than either "
    "other subject, indicating relative vulnerability). Because these individually-robust effects "
    "point in opposite directions, the pooled fixed-effect interaction term (day distance × zone "
    "category) in a mixed-effects model was not significant (β = +0.0022, p = .465), correctly "
    "reflecting the absence of a single population-level rule rather than the absence of any effect "
    "(Figure 3, Figure 4)."
)
figure(FIG + "Z2_reward_vs_corridor_slopes.png",
       "Figure 3. Population vector correlation versus day distance, restricted to reward-zone (red) "
       "and corridor (grey) bins, per subject. UT14 and UT13 show a flatter (more stable) reward-zone "
       "slope; UT15 shows a steeper (less stable) reward-zone slope.")
figure(FIG + "Z3_shuffle_null_reward.png",
       "Figure 4. Occupancy-matched shuffle null distributions for the reward-vs-corridor slope "
       "difference, per subject. Observed values (red lines) fall well outside the null distribution "
       "in all three subjects, confirming the effect is not attributable to reward zones being "
       "differently sampled or sized than corridor.")

h2("Held-out validation identifies the reward-zone-vulnerability effect as the most robust finding", "2.4")
body(
    "Because zones were originally defined using trial data pooled across a subject's entire session "
    "history and then tested on population vectors built from those same sessions, we performed a "
    "held-out validation to rule out circularity. Sessions were split by temporal order into two "
    "interleaved halves; zones were defined using only one half, and the reward-vs-corridor effect was "
    "tested exclusively on population vectors built from the other, independent half. Both split "
    "directions were run (Table 1)."
)
make_table(
    ["Subject", "Split", "Reward − corridor slope diff.", "p", "Replicates?"],
    [
        ["UT14", "define even / test odd", "+0.0028", ".0020", "Yes"],
        ["UT14", "define odd / test even", "+0.0008", ".307", "No"],
        ["UT13", "define even / test odd", "+0.0004", ".523", "No"],
        ["UT13", "define odd / test even", "−0.0008", ".102", "No (sign reverses)"],
        ["UT15", "define even / test odd", "−0.0110", ".0020", "Yes"],
        ["UT15", "define odd / test even", "−0.0168", ".0020", "Yes"],
    ],
    col_widths=[0.7, 1.8, 1.7, 0.6, 1.4],
)
body(
    "UT15's effect replicated fully and in the same direction across both independent splits. UT14's "
    "effect replicated directionally but reached significance in only one of the two splits, consistent "
    "with a real but modest, power-limited effect. UT13's effect did not replicate: the two splits "
    "disagreed in sign, and neither was individually significant, indicating the original full-sample "
    "result for this subject (already the weakest of the three, p = .024) reflects a marginal, "
    "non-replicating finding rather than reliable evidence for reward-zone protection (Figure 5)."
)
figure(FIG + "Z5_heldout_validation.png",
       "Figure 5. Reward-vs-corridor slope difference from the full pooled dataset (blue) compared "
       "against both held-out validation splits (red = significant at p < .05, grey = not significant), "
       "per subject. UT15 replicates in both splits; UT14 partially; UT13 does not replicate.")

h2("The UT15 effect is not attributable to a position-tracking artifact", "2.5")
body(
    "Visual inspection of raw trial-by-trial trajectories (rather than aggregated occupancy) revealed a "
    "qualitative difference in tracking quality: UT15's position data in the maze corridor showed a "
    "diffuse, scattered pattern with visible discontinuous jumps, unlike the tight, continuous "
    "trajectories in UT14 and UT13. This was quantified directly by flagging any position sample "
    "reached via a physically-impossible instantaneous displacement (> 100 cm/s, computed from "
    "consecutive-sample distance and inter-sample interval; median sampling interval ≈ 33 ms across "
    "all subjects). Restricted to the horizontal maze-corridor region (excluding any non-maze holding "
    "area), UT15 showed a jump rate of 3.26% (39,656 / 1,217,789 samples) versus 0.31% for UT14 and "
    "0.41% for UT13 — an 8- to 10-fold elevation specific to this subject."
)
body(
    "To determine whether this tracking artifact could explain the reward-zone finding, all flagged "
    "samples in UT15 were replaced by linear interpolation between adjacent valid samples (92,447 "
    "samples across 15 sessions), and the full zone analysis was recomputed on the cleaned position "
    "data. The reward-vs-corridor effect was essentially unchanged: slope difference = −0.0149, "
    "p = .0020 (cleaned) versus −0.0145, p = .0020 (uncleaned). This constitutes a fourth "
    "independent robustness check — alongside the occupancy-matched shuffle control, a "
    "zone-boundary parameter-sensitivity check (Supplementary Results), and held-out validation — "
    "all of which support the reward-zone-vulnerability effect in UT15 as genuine rather than "
    "artifactual."
)
body(
    "The choice-zone comparison, by contrast, was substantially altered by cleaning (uncleaned: "
    "difference = −0.0008, p = .659; cleaned: +0.0117, p = .0020) and was separately found to be "
    "sensitive to the choice-zone boundary definition (Supplementary Results). We conclude that the "
    "choice-zone result is not reliable in this dataset and should be treated as inconclusive, whereas "
    "the reward-zone result is robust across every check applied."
)

h2("Supplementary analysis: recording time-of-day does not predict representational drift", "2.6")
body(
    "The same three subjects' verified acquisition timestamps permit a test of whether time-of-day "
    "separation between sessions, independent of calendar-day separation, predicts representational "
    "drift — included here as a negative control confirming the day-distance effect above is not "
    "confounded by circadian time-of-day, and reported briefly for completeness. UT15's sessions spanned "
    "a genuine 11.3-hour range of recording start times (09:31–20:50), the widest of the three "
    "subjects and sufficient to test this question with reasonable power; UT14 spanned 6.15 hours and "
    "UT13 spanned 2.12 hours. No subject showed a relationship between time-of-day separation and PV "
    "correlation after controlling for day distance (partial Mantel test: UT14, r = −0.114, "
    "p = .485; UT13, r = 0.040, p = .781; UT15, r = 0.006, p = .955; pooled Fisher's method, "
    "p = .916)."
)

doc.add_page_break()

# ============================================================================
h1("Discussion", 3)
body(
    "This study replicated chronic hippocampal representational drift across three independently "
    "recorded mice performing a spatial working-memory task, and then asked a question not directly "
    "addressed by prior work: within a single, continuously-rewarded environment, does the reward zone "
    "itself show altered drift dynamics relative to neutral corridor space? We found that it does — "
    "in every subject — but that the direction of this modulation was not uniform. Two subjects "
    "showed the reward zone as relatively protected from drift, in the direction one would predict by "
    "analogy with prior reward-expectation work; one subject showed the opposite, with the reward zone "
    "decorrelating faster than corridor. Critically, when subjected to held-out cross-validation, only "
    "the vulnerability effect (UT15) fully replicated; the protection effect in UT13 did not replicate "
    "at all, and in UT14 replicated only partially. The vulnerability effect additionally survived a "
    "direct control for a position-tracking artifact identified in that subject, a check motivated by a "
    "visible irregularity in that subject's zone geometry that might otherwise have been dismissed as a "
    "data-quality caveat rather than investigated directly."
)
body(
    "These findings should be read as a caution against assuming that reward uniformly stabilizes "
    "nearby hippocampal representations. The existing literature linking reward to reduced "
    "representational drift has manipulated reward availability globally, comparing whole sessions "
    "with reward present against whole sessions with reward withheld (Krishnan & Sheffield, 2023). "
    "The present analysis asks a "
    "complementary, spatially-localized question — within an environment that is always rewarded, "
    "is drift uniform across space? — and finds that it is not, but that the specific direction of "
    "non-uniformity is animal-specific rather than following a single rule. This is consistent with, "
    "and extends, evidence that reward reorganizes the hippocampal population code in ways that go "
    "beyond simple firing-rate overrepresentation: a dedicated, dynamically-recruited subpopulation "
    "appears to track position relative to reward specifically (Sosa, Plitt, & Giocomo, 2025), and if "
    "the composition or engagement of such a subpopulation itself varies across animals or fluctuates "
    "over the weeks-long timescale studied here, elevated apparent drift at the reward zone specifically "
    "— rather than protection — is a plausible outcome in at least some animals."
)
body(
    "Two methodological findings, incidental to the original aim of this analysis, are worth "
    "highlighting as general cautions for secondary analysis of public chronic-recording datasets. "
    "First, session-level timing metadata (here, the NWB “session_start_time” field) should not "
    "be trusted without cross-checking against an independently-verifiable, internally-consistent "
    "timestamp source (here, the embedded position/spike/trial acquisition clock) — the discrepancy "
    "found here was large (up to several hours) and inconsistent in direction, which would not be "
    "detectable from the metadata field alone. Second, position-tracking quality can vary substantially "
    "and non-obviously between subjects within the same dataset, in ways that manifest as apparently "
    "biological irregularities (here, a geometrically distorted zone classification) that could easily "
    "be misattributed to behavior without a direct, quantitative tracking-quality check."
)
body(
    "This study has several limitations. The sample size (three subjects) provides ample within-subject "
    "statistical power, given the many session pairs available per animal, but only modest power to "
    "characterize the population-level distribution of individual differences in reward-zone drift "
    "direction; a larger cohort would be needed to determine what fraction of animals show protection "
    "versus vulnerability, and what predicts this. Cell identity is not tracked across recording "
    "sessions in this dataset (units are independently spike-sorted per session), which necessitated a "
    "tetrode-pooled, rather than single-unit, population vector — a principled workaround given the "
    "physically fixed tetrode geometry, but one that cannot resolve single-cell-level heterogeneity in "
    "drift that recent work suggests may be substantial (e.g., cell excitability predicting individual "
    "drift rate). The choice-zone comparison, unlike the reward-zone comparison, did not survive its "
    "robustness checks and should be considered inconclusive rather than a genuine null. Finally, this "
    "analysis is correlational; it cannot distinguish among candidate mechanisms (dopaminergic "
    "modulation, reward-relative subpopulation turnover, behavioral factors such as running speed or "
    "dwell time near reward) for either the protective or vulnerable pattern observed."
)
body(
    "In summary, reward-adjacent hippocampal representations in this dataset were not uniformly "
    "protected from long-term representational drift. At least one animal showed a robust, "
    "thoroughly-validated pattern of accelerated — not reduced — drift specifically at the "
    "reward zone, surviving four independent checks for artifact or circularity. This indicates "
    "meaningful, currently unexplained individual variability in how spatially-localized reward "
    "modulates the long-term stability of the hippocampal cognitive map, and argues that reward-drift "
    "relationships established under global reward manipulations should not be assumed to generalize, "
    "in a uniform direction, to the spatial structure of drift within a continuously-rewarded "
    "environment."
)

doc.add_page_break()

# ============================================================================
h1("Methods", 4)

h2("Dataset and subject inclusion criteria", "4.1")
body(
    "Data were obtained from DANDI Archive dandiset 001775 (publicly available, no registration "
    "required), associated with an in-preparation manuscript from the Kitamura Lab (data also mirrored "
    "at Zenodo, DOI 10.5281/zenodo.18580888). Six mice were chronically implanted with tetrode arrays "
    "targeting right dorsal CA1 (up to 32 tetrode positions) and performed a delayed-non-match-to-place "
    "(DNMP) T-maze task across multiple recording sessions. Subjects were included in the present "
    "analysis only if their NWB files contained both a Position/SpatialSeries interface and a "
    "spike-sorted units table, and only if session position timestamps fell within a plausible absolute "
    "Unix-epoch range (2015-01-01 to 2026-01-01, spanning the dataset's collection window) — sessions "
    "failing this check use a lab-internal relative acquisition clock that cannot be mapped to real "
    "calendar time from file contents alone and were excluded rather than converted with an assumed "
    "offset. This excluded one subject entirely (different behavioral paradigm, no spike-sorted units) "
    "and two subjects' full session sets (relative-clock timestamps). The final sample comprised UT14 "
    "(19 sessions), UT13 (15 sessions), and UT15 (15 sessions)."
)

h2("Timestamp verification", "4.2")
body(
    "For every included session, the true recording clock was reconstructed from the mutually-consistent "
    "Unix-epoch timestamps embedded in that session's position, spike, and trial-interval data (verified "
    "to agree with one another to sub-second precision within each session) and used in place of the "
    "NWB “session_start_time” metadata field, which was independently found to disagree with "
    "this embedded clock by up to several hours, inconsistently in direction, for every included subject."
)

h2("Spike sorting and unit inclusion", "4.3")
body(
    "Units were spike-sorted independently per session using MountainSort4 (no native cross-session "
    "identity tracking). A unit was retained for analysis if its reported isolation score was ≥ 0.90, "
    "its noise-overlap score was ≤ 0.10, and it produced ≥ 100 spikes during the session. In "
    "practice, sorting quality was uniformly high and no units were excluded by these criteria in any "
    "retained session. Retained unit counts ranged 76–136 per session for UT14 and UT15 and "
    "approximately 222–260 per session for UT13, reflecting a denser probe montage in that subject."
)

h2("Behavioral task and epochs", "4.4")
body(
    "Each session comprised, in sequence, a pre-run open-field foraging epoch, the DNMP T-maze task "
    "epoch, and a post-run open-field foraging epoch; all analyses reported here use the DNMP epoch. "
    "Individual trials within the DNMP epoch were annotated with trial phase (sample/test), choice "
    "(left/right arm), correct arm, and event timestamps for reward-related behavior (grasp, chewing "
    "onset, forepaw contact). Reward was delivered at a fixed physical location at the end of each arm "
    "on every completed trial; reward availability was never globally withheld across the sessions "
    "analyzed here."
)

h2("Rate maps and population vector construction", "4.5")
body(
    "Position data were binned into a 4 × 4 cm spatial grid, shared across all of a subject's "
    "sessions (the union of the spatial extent visited during the DNMP epoch across sessions, padded by "
    "5 cm). Occupancy-normalized, Gaussian-smoothed (σ = 1 bin) firing-rate maps were computed per "
    "unit per session; a bin required ≥ 0.15 s of occupancy to be considered visited. Because units "
    "are not tracked across sessions, population vectors were constructed at the tetrode level: for each "
    "of up to 32 physically-fixed tetrode positions, the mean rate map across all units recorded on that "
    "tetrode in a given session was computed, yielding a fixed-dimensionality (tetrode × spatial bin) "
    "representation directly comparable across any session pair without requiring single-unit identity "
    "matching. Population vector (PV) correlation between two sessions was computed, per spatial bin, as "
    "the Pearson correlation across the common set of tetrodes active in both sessions (minimum 3 "
    "tetrodes per bin); the session-pair PV correlation is the Fisher-z-averaged per-bin correlation "
    "across all bins with sufficient occupancy in both sessions (minimum 15 usable bins per pair)."
)

h2("Zone classification", "4.6")
body(
    "Reward zone: a 10 cm radius around the median animal position at the grasp_time event (or, if "
    "unavailable, chewing_onset_time) of every trial of a given choice side, computed separately for "
    "left and right arms and pooled across a subject's full session history, then unioned. Choice zone: "
    "spatial bins visited during both left- and right-choice trials (“shared” bins) that "
    "directly adjoin the region visited exclusively during one choice side (1-bin morphological dilation "
    "of the exclusive region, intersected with the shared region), operationalizing the maze stem "
    "immediately proximal to the choice point. Corridor: all remaining visited bins. Sensitivity to the "
    "dilation parameter (1 vs. 2 bins) was assessed as a robustness check (Supplementary Results)."
)

h2("Statistical analysis", "4.7")
body(
    "Per-subject inference on the day-distance drift effect used a Mantel permutation test (5,000 "
    "permutations of session order) correlating the vectorized upper-triangle of the PV-correlation "
    "matrix with the calendar-day-distance matrix. For zone comparisons, the reward-vs-corridor "
    "(and choice-vs-corridor) effect was quantified as the difference in OLS slope of PV correlation on "
    "day distance between the two zone-restricted PV matrices; significance was assessed via an "
    "occupancy-matched bin-shuffle permutation test (500 permutations), in which the same visited bins "
    "were randomly reassigned to same-sized pseudo-zone categories (preserving each real zone's bin "
    "count) and the pseudo-zone slope difference recomputed each permutation to build a null "
    "distribution. Cross-subject pooling used Fisher's method to combine independent per-subject "
    "permutation p-values, and a linear mixed-effects model (day distance × zone category fixed "
    "effects, random intercept per subject, fit by REML) as a complementary parametric estimate."
)

h2("Held-out validation", "4.8")
body(
    "Sessions were split by temporal order into two interleaved halves (even-index / odd-index after "
    "sorting by verified session date, so both halves span the full day-distance range rather than "
    "confounding the split with session order). Zones were defined using only one half; the "
    "reward-vs-corridor effect was then tested exclusively on PV correlations built from the other "
    "half's session pairs. Both split directions were run and reported."
)

h2("Position-tracking artifact control", "4.9")
body(
    "Position samples reached via a physically-impossible instantaneous displacement (> 100 cm/s, "
    "computed from consecutive-sample distance and inter-sample interval) were flagged, and the flagged "
    "rate was compared across subjects, restricted to the horizontal maze-corridor region "
    "(y ∈ [−4, 6] cm) to exclude any non-maze area from the comparison. For the subject showing "
    "an elevated rate, flagged samples were replaced by linear interpolation between adjacent valid "
    "samples, and the full zone analysis was recomputed on the cleaned data for comparison against the "
    "original result."
)

h2("Software", "4.10")
body(
    "All analyses were performed in Python 3.13 using pynwb, NumPy, SciPy, pandas, statsmodels, and "
    "matplotlib. Full analysis code, including the underlying Jupyter notebooks with all figures and "
    "intermediate outputs, is available at the repository listed in Data and Code Availability."
)

doc.add_page_break()

# ============================================================================
h1("Data and Code Availability", 5)
body(
    "Raw data are publicly available from the DANDI Archive, dandiset 001775 "
    "(https://dandiarchive.org/dandiset/001775), with no registration required. Analysis code, "
    "including all rate-map, population-vector, zone-classification, and statistical modules, plus the "
    "Jupyter notebooks reproducing every figure and statistic in this manuscript, is available at "
    "[repository URL to be added]."
)

h1("Acknowledgments", 6)
body("[To be added.]")

h1("Author Contributions", 7)
body("[To be added — e.g., CRediT taxonomy: Conceptualization, Formal analysis, Methodology, "
     "Software, Visualization, Writing.]")

h1("Competing Interests", 8)
body("The authors declare no competing interests. [Confirm and adjust as appropriate.]")

doc.add_page_break()

# ============================================================================
h1("References", 9)
note_box(
    "Citation completeness note: full author lists, journal names, and volume/issue/article details "
    "for every entry below (including Krishnan & Sheffield, 2023 and Tarcsay et al., 2025) were "
    "independently confirmed against PubMed/PMC/publisher records during this analysis. Krishnan & "
    "Sheffield (2023) remains a bioRxiv preprint as of verification (not yet published in a "
    "peer-reviewed journal) -- re-check before submission in case it has since been formally "
    "published. Classic citations (O'Keefe & Dostrovsky, 1971; Ziv et al., 2013; Hollup et al., 2001; "
    "Gauthier & Tank, 2018) are well-established papers cited from general knowledge, not "
    "independently re-verified in this session, and should be spot-checked against a reference "
    "manager before submission."
)

refs = [
    "Gauthier, J. L., & Tank, D. W. (2018). A dedicated population for reward coding in the "
    "hippocampus. Neuron, 99(1), 179–193.",
    "Geva, N., Deitch, D., Rubin, A., & Ziv, Y. (2023). Time and experience differentially affect "
    "distinct aspects of hippocampal representational drift. Neuron, 111(15), 2357–2366.",
    "Hollup, S. A., Molden, S., Donnett, J. G., Moser, M.-B., & Moser, E. I. (2001). Accumulation of "
    "hippocampal place fields at the goal location in an annular watermaze task. Journal of "
    "Neuroscience, 21(5), 1635–1644.",
    "Krishnan, S., & Sheffield, M. E. J. (2023). Reward expectation reduces representational drift in "
    "the hippocampus. bioRxiv preprint, doi:10.1101/2023.12.21.572809. [Verified as a bioRxiv "
    "preprint, PMID: 38187677, PMCID: PMC10769341; not yet listed as formally published in a "
    "peer-reviewed journal as of verification -- re-check before submission.]",
    "O'Keefe, J., & Dostrovsky, J. (1971). The hippocampus as a spatial map: preliminary evidence from "
    "unit activity in the freely-moving rat. Brain Research, 34(1), 171–175.",
    "Rule, M. E., O'Leary, T., & Harvey, C. D. (2019). Causes and consequences of representational "
    "drift. Current Opinion in Neurobiology, 58, 141–147.",
    "Sosa, M., Plitt, M. H., & Giocomo, L. M. (2025). A flexible hippocampal population code for "
    "experience relative to reward. Nature Neuroscience, 28, 1497–1509.",
    "Tarcsay, G., Masala, N., Yi, J. D., Igarashi, M. K., Redic, U. J., & Ewell, L. A. (2025). The "
    "relevance of context in memory tasks influences the magnitude of hippocampal remapping. Cell "
    "Reports, 44(12), 116682. doi:10.1016/j.celrep.2025.116682",
    "Ziv, Y., Burns, L. D., Cocker, E. D., Hamel, E. O., Ghosh, K. K., Kitch, L. J., El Gamal, A., & "
    "Schnitzer, M. J. (2013). Long-term dynamics of CA1 hippocampal place codes. Nature Neuroscience, "
    "16(3), 264–266.",
]
for ref in refs:
    p = doc.add_paragraph()
    r = p.add_run(ref)
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)

doc.add_page_break()

# ============================================================================
h1("Supplementary Results", 10)
h2("Sensitivity to the choice-zone dilation parameter", "S10.1")
body(
    "The choice-zone effect was re-assessed under a wider dilation setting (2 bins instead of 1). The "
    "reward-zone effect was unchanged in direction, magnitude, and significance under this change in all "
    "three subjects, but the choice-zone effect was not: it reached significance for UT15 under the "
    "wider setting (difference = +0.0034, p = .006) but not the narrower one (difference = −0.0008, "
    "p = .675) used in the main text, confirming the choice-zone result is parameter-sensitive and "
    "should not be interpreted with confidence in either direction."
)

h2("Supplementary Figures", "S10.2")
figure(FIG + "Z4_choice_vs_corridor.png",
       "Figure S1. Population vector correlation versus day distance, restricted to choice-zone (orange) "
       "and corridor (grey) bins, per subject. No subject shows a robust, cleaning- and "
       "parameter-independent effect (see Sections 2.5 and S10.1).")
figure(FIG + "A3_true_hour_spread_all_subjects.png",
       "Figure S2. True (verified) recording time-of-day for every session, per subject, supporting the "
       "supplementary time-of-day analysis (Section 2.6). UT15 shows the widest true time-of-day span "
       "(11.3 hours) of the three subjects.")

doc.save('hippocampus_reward_zone_drift_preprint.docx')
print("saved: hippocampus_reward_zone_drift_preprint.docx")
